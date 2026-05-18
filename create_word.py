from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_guideline():
    doc = Document()
    
    # Title
    title = doc.add_heading('PANDUAN INTEGRASI FRONTEND - THRIFTLY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Versi: 2.0 (Revisi Core API & Advanced Verification)')
    doc.add_paragraph('Status API: Production Ready on VPS')
    
    # Section 1
    doc.add_heading('1. PEMBAYARAN (MIDTRANS CORE API - CUSTOM UI)', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('Metode ini menggunakan UI kustom sepenuhnya (tanpa Snap Popup).').bold = True
    doc.add_paragraph('• Alur: User pilih metode -> Frontend tembak API -> Backend return data VA/QR/Code -> Frontend render ke layar.')
    doc.add_paragraph('• UI Requirements: Tampilkan nomor Virtual Account, tombol "Salin", dan instruksi pembayaran.')
    doc.add_paragraph('• Status: Tambahkan tombol "Cek Status Pembayaran" untuk memicu refresh data transaksi.')
    
    # Section 2
    doc.add_heading('2. VERIFIKASI IDENTITAS (KTP)', level=1)
    doc.add_paragraph('• Submission: Gunakan FormData untuk mengirim NIK, Nama, dan Foto KTP.')
    doc.add_paragraph('• Rejection Logic: Jika status "rejected", tampilkan alasan penolakan (ktp_rejection_reason) dan aktifkan kembali tombol upload.')
    doc.add_paragraph('• Verified Status: Tampilkan badge hijau centang jika "is_ktp_verified" bernilai true.')
    
    # Section 3
    doc.add_heading('3. VERIFIKASI EMAIL (LINK KONFIRMASI)', level=1)
    doc.add_paragraph('• Mekanisme: Satu klik melalui email (tanpa kode OTP).')
    doc.add_paragraph('• Frontend Action: Tombol "Kirim Link Verifikasi" menembak POST /api/email/verification-notification.')
    doc.add_paragraph('• Feedback: Tampilkan pesan sukses bahwa link telah dikirim ke Gmail user.')
    
    # Section 4
    doc.add_heading('4. VERIFIKASI WHATSAPP (OTP)', level=1)
    doc.add_paragraph('• Request: Tembak POST /api/otp/send untuk mengirim kode ke WA user.')
    doc.add_paragraph('• Verify: Sediakan input 6 digit kode, kirim ke POST /api/otp/verify.')
    
    # Section 5
    doc.add_heading('5. SISTEM NOTIFIKASI (BELL ICON)', level=1)
    doc.add_paragraph('• Bell Icon: Tambahkan di Header untuk notifikasi real-time.')
    doc.add_paragraph('• Alerts: Notifikasi harus muncul jika verifikasi KTP ditolak atau pembayaran dikonfirmasi.')
    
    # Section 6
    doc.add_heading('6. ADMIN PANEL (REVIEW KTP)', level=1)
    doc.add_paragraph('• List: Tabel khusus menampilkan user dengan status pending.')
    doc.add_paragraph('• Action: Tombol Approve (Selesai) dan Tombol Reject (Wajib mengisi alasan penolakan).')
    
    # Save the document
    doc.save('Panduan_Integrasi_Frontend_Thriftly.docx')
    print("Document created successfully: Panduan_Integrasi_Frontend_Thriftly.docx")

if __name__ == "__main__":
    create_guideline()
