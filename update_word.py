from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 102, 204) # Blueish for code

def update_guideline():
    doc = Document()
    
    # Title
    title = doc.add_heading('DOKUMEN TEKNIS INTEGRASI FRONTEND - THRIFTLY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Dokumen ini berisi langkah-langkah implementasi fitur Verifikasi & Pembayaran untuk Developer Frontend.')

    # SECTION 1: MIDTRANS CORE API
    doc.add_heading('1. INTEGRASI PEMBAYARAN (MIDTRANS CORE API)', level=1)
    doc.add_paragraph('Gunakan API langsung untuk merender UI pembayaran kustom Anda sendiri.')
    doc.add_paragraph('Contoh penanganan data VA (Virtual Account) dari Backend:')
    add_code_block(doc, """// Contoh Response dari API Backend:
// { "va_numbers": [{ "bank": "bca", "va_number": "12345678" }], "status": "pending" }

const renderPaymentInfo = (paymentData) => {
    return (
        <div className="p-4 border rounded-xl bg-gray-50">
            <h4 className="font-bold">Nomor Virtual Account {paymentData.va_numbers[0].bank.toUpperCase()}</h4>
            <div className="flex gap-2 items-center text-2xl font-mono py-2">
                {paymentData.va_numbers[0].va_number}
                <button onClick={() => copyToClipboard(paymentData.va_numbers[0].va_number)}>Salin</button>
            </div>
            <p className="text-sm text-gray-500">Silakan lakukan transfer sebelum 24 jam.</p>
        </div>
    );
};""")

    # SECTION 2: KTP VERIFICATION LOGIC
    doc.add_heading('2. LOGIKA VERIFIKASI IDENTITAS (KTP)', level=1)
    doc.add_paragraph('Handle tiga kondisi status: null/empty, pending, verified, dan rejected.')
    doc.add_paragraph('Implementasi pengiriman data menggunakan FormData:')
    add_code_block(doc, """const handleSubmitKtp = async () => {
    const formData = new FormData();
    formData.append('ktp_nik', nik);
    formData.append('ktp_name', name);
    formData.append('ktp_image', selectedFile); // File object dari input type="file"

    try {
        await axios.post('/api/user/verify-ktp', formData, {
            headers: { 'Content-Type': 'multipart/form-data' }
        });
        toast.success("Data berhasil dikirim, menunggu peninjauan admin.");
    } catch (err) {
        toast.error("Gagal mengirim data.");
    }
};""")

    doc.add_paragraph('Jika status === "rejected", tampilkan alasan penolakan agar user bisa kirim ulang:')
    add_code_block(doc, """{user.ktp_status === 'rejected' && (
    <div className="bg-red-50 p-4 border border-red-200 rounded-lg mb-4">
        <p className="text-red-700 font-bold">Verifikasi Ditolak</p>
        <p className="text-red-600 text-sm">Alasan: {user.ktp_rejection_reason}</p>
        <p className="text-xs mt-2 italic">Silakan upload kembali foto KTP yang lebih jelas.</p>
    </div>
)}""")

    # SECTION 3: EMAIL VERIFICATION LINK
    doc.add_heading('3. VERIFIKASI EMAIL (LINK KONFIRMASI)', level=1)
    doc.add_paragraph('Cukup panggil endpoint untuk mentrigger pengiriman email.')
    add_code_block(doc, """const sendVerificationLink = async () => {
    setLoading(true);
    await axios.post('/api/email/verification-notification');
    setLoading(false);
    toast.success("Link verifikasi telah dikirim ke Gmail Anda!");
};""")

    # SECTION 4: WHATSAPP OTP (FONNTE)
    doc.add_heading('4. VERIFIKASI WHATSAPP (OTP)', level=1)
    doc.add_paragraph('Alur dua langkah: Kirim kode -> Verifikasi kode.')
    add_code_block(doc, """// 1. Kirim Kode
await axios.post('/api/otp/send', { phone: user.no_telp });

// 2. Verifikasi Kode
await axios.post('/api/otp/verify', { 
    phone: user.no_telp, 
    code: otpInput 
});""")

    # SECTION 5: NOTIFIKASI LONCENG
    doc.add_heading('5. NOTIFIKASI REAL-TIME (HEADER BELL)', level=1)
    doc.add_paragraph('Gunakan data dari Auth Context untuk mendeteksi update status secara global.')
    doc.add_paragraph('Contoh indikator notifikasi baru di Header:')
    add_code_block(doc, """<div className="relative">
    <BellIcon className="w-6 h-6" />
    {user.ktp_status === 'rejected' && (
        <span className="absolute top-0 right-0 h-3 w-3 bg-red-500 rounded-full border-2 border-white animate-pulse" />
    )}
</div>""")

    # Save
    doc.save('Panduan_Integrasi_Frontend_Thriftly_V2.docx')
    print("Document Updated successfully: Panduan_Integrasi_Frontend_Thriftly_V2.docx")

if __name__ == "__main__":
    update_guideline()
